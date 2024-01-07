import ast
import os
import openai
from docx2txt import process
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Input and Output Folder Global Definitions
cv_folder = "CVs"
output_folder = "Formatted CVs"

def cv_reader(unformatted_cv):
    # Set OpenAI API key & Initialize GPT Client
    openai.api_key = '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'

    # Initialize AI's Personality
    conversation = [
    {"role": "system", "content": """I will pass you a CV template divided into sections with [SECTION xyz] subheaders. I want you to read the Instructions mentioned at the bottom of the template such that I can pass you another actual CV and you can format it according to the template and give me back a list as specified in the instructions. Do NOT invent your own sections, regardless of whether the CV i give you has extra content, stick to the 5 sections mentioned in the template. I want you to pass me a LIST object suitable for interpretation in Python:

[SECTION 1A] NAME
[SECTION 1B] Location | LinkedIn| Phone | Email
[SECTION 2] EDUCATION
University                                                                                                                                         Location
Major/Course                                                                                                                                          Dates                                   
•	GPA[if 3.2 or above], Organizations, Coursework, etc.                                                                                                                                                   

[SECTION 3] WORK EXPERIENCE 
Company                                                                                                                                          Location
Position                                                                                                                                                   Dates
•	This section regarding experiences has bulleted accomplishments, which provide examples of when you successfully used the skills employers are seeking. Make sure you have between 2 and 6 bullet points in each section. 
•	Your bullet points should start with a strong action verb, which then follows with an explanation of what you were doing, describing how you did it, and most importantly if applicable, any achievements. Statements should convey your strengths/proficiencies in one or more skills that intrigue the employer by showing examples of when you have used them. 
•	When writing about your experience, consider these questions: What was the result/outcome of your work? What were your accomplishments? How did you impact the organization? What skills/knowledge did you grow? How does this experience relate to your internship/employment goal?
Company                                                                                                                                                                        Location
Position                                                                                                                                                                                    Dates
•	Your bullet statements should be in proper tense, using-ed for past experiences and present tenses for current positions. Make sure that your writing is free of grammatical errors and punctuation. · 
•	When including numerical achievements during your experiences, make sure to include (if applicable) the quantity, population, frequency, and impact of your work whenever possible. 
•	To make your resume flow, read it over. Check and see if it is easy to read with no overflowing of text. You should avoid the usage of different colors, multiple fonts, pictures, and brief/too dense information. Your resume should show who you are while being professional.

[SECTION 4] LEADERSHIP EXPERIENCE (OPTIONAL)
Company                                                                                                                                          Location
Position                                                                                                                                                   Dates
•	This section is optional if you have various leadership experiences and other activities you want employers to know. By having multiple sections, it allows you to emphasize your most relevant experience. 
•	Positions within this section should be formatted similarly to previous experience sections, including bullet points if necessary. 
•	You may also include work experiences that may not be directly related to the job/internship you are applying to but add to your credibility by exemplifying your past work experiences.

[SECTION 5] SKILLS & ADDITIONAL TRAINING
Skills: These skills should be concrete and testable. These should not be soft skills like communication, organizational, and interpersonal skills, but instead incorporated into your bulleted accomplishment statements above. Interests: What type of additional training did you perform, or certificates did achieve? 

Instructions:

Return a list [SECTION 1A, SECTION 1B, SECTION 2, SECTION 3, SECTION 4, SECTION 5]

Where:
SECTION 1A is Name

SECTION 1B should be a list [Location, LinkedIn, Phone, Email]. If one of these is missing, just input a placeholder for example LOCATION.

SECTION 2 should be a list of education [number of sublists, sublist1, sublist2,…] . The first index represents how many entries/sublists there are and the following indices contain those work experiences in sublists. Each Education sublist should be of the form [University/School Name, Location, Major/Course, Dates, [a list containing GPA, organizations, coursework, etc]]. If there are none, make the first index 0.

SECTION 3 should be a list of work experiences [number of sublists, sublist1, sublist2,…]. The first index represents how many entries/sublists there are and the following indices contain those work experiences in sublists. Each Leadership experience sublist should be of the form [Company, Location, Position, Dates, Description of experience]

SECTION 4 should be a list of leadership experiences [number of sublists, sublist1, sublist2,…]. The first index represents how many entries/sublists there are and the following indices contain those leadership experiences in sublists. Each Leadership experience sublist should be of the form [Company, Location, Position, Dates, Description of experience]

SECTION 5 should be a list of skills and additional training [number of entires, entry1, entry2,…] where the first index represents how many entries there are and the following indices contain those additional skill and additional training bullet points. If there are none, make the first index 0.
     
Remember, the resultant formatted list should ONLY have these 5 sections and NO MORE, no additional references or interests section for example. NOTHING else.

Ensure University/Company Name, Location, Dates, Degree/Position, Dates, Description of Entry are 5 seperate entries in each of the sublists in sections 2 - education, 3 - work experience, and 4 - leadership experience. If nothing can be found, just leave '' for that empty section.

Since the result will be used in code directly, please put any comments you have as a 6th "section" in the list object directly. Ensure that the size of the resultant list is only these 6 sections and the additional comments section [SECTION 1A, SECTION 1B, SECTION 2, SECTION 3, SECTION 4, SECTION 5].
     
Lastly, remember, the number of sublists/entries in each section should be the first index in that respective list, not its own entry (as the indexing to access the information depends on this consistency). Likewise, pleasure leave no code comments in between the list as well and return NOTHING but the list. No additional text outside the list.
     """},
    ]
    conversation.append({"role": "assistant", "content": """[
    "John Doe",  
    
    ["Location", "LinkedIn", "Phone", "Email"],  
    
    [
        2,  
        ["University", "Location", "Major/Course", "Dates", ["GPA", "Organizations", "Coursework"]],
        ["University", "Location", "Major/Course", "Dates", ["GPA", "Organizations", "Coursework"]]
    ],  
    
    [
        2,  
        ["Company", "Location", "Position", "Dates", "Description of experience"],
        ["Company", "Location", "Position", "Dates", "Description of experience"]
    ], 
    
    [
        1, 
        ["Company", "Location", "Position", "Dates", "Description of experience"]
    ],  
    
    [
        2,  
        "Concrete and Testable Skill",
        "Additional Training or Certificates Achieved"
    ],  
    
    # Comments:
    # - Ensure the consistency of sublist structure across all sections for accurate indexing.
    # - The first index of each section represents the number of entries/sublists.
    # - Follow the specified format strictly to facilitate proper interpretation in Python.
    
]"""})

    # Send in New CV
    conversation.append({"role": "user", "content": unformatted_cv})

    # Get AI-generated response
    formatted_cv_list = chat_with_model(conversation)
    conversation.append({"role": "assistant", "content": formatted_cv_list})
    print(formatted_cv_list)
    # print(formatted_cv_list)
    return ast.literal_eval(formatted_cv_list)

# Helper Function to Create a conversation and get AI-generated responses
def chat_with_model(messages):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    return response.choices[0].message["content"]

def docx_writer(formatted_cv_list):
    # Takes a Formatted List and Outputs Desired CV as per Template
    # List [SECTION 1A, SECTION 1B, SECTION 2, SECTION 3, SECTION 4, SECTION 5]

    # Where:
    # SECTION 1A is Name
    # SECTION 1B should be a list [Location, LinkedIn, Phone, Email]. If one of these is missing, just input a placeholder for example LOCATION.
    # SECTION 2 should be a list of education. The first index represents how many entries there are and the following indices contain those work experiences in sublists. Each Education sublis should be of the form [University/School Name, Location, Major/Course, Dates, [a list containing GPA, organizations, coursework, etc]]. If there are none, make the first index 0.
    # SECTION 3 should be a list of work experiences. The first index represents how many entries there are and the following indices contain those work experiences in sublists. Each Leadership experience sublist should be of the form [Company, Location, Position, Dates, Description of experience]
    # SECTION 4 should be a list of leadership experiences. The first index represents how many entries there are and the following indices contain those leadership experiences in sublists. Each Leadership experience sublist should be of the form [Company, Location, Position, Dates, Description of experience]
    # SECTION 5 should be a list of skills and additional training where the first index represents how many entries there are and the following indices contain those additional skill and additional training bullet points. If there are none, make the first index 0.
    name_header = formatted_cv_list[0]
    print(name_header)
    desc_header = formatted_cv_list[1]
    print(desc_header)

    # Extract Relevant Sectioned Information from the List
    if formatted_cv_list[2] == 0:
        education_section_sublists = "ERROR"
    else:
        education_section_sublists = formatted_cv_list[2][1:]
        print(education_section_sublists)

    if formatted_cv_list[3] == 0:
        work_experience_sublists = "ERROR"
    else:
        work_experience_sublists = formatted_cv_list[3][1:]
        print(work_experience_sublists)

    if formatted_cv_list[4] == 0:
        leadership_experience_sublists = "ERROR"
    else:
        leadership_experience_sublists = formatted_cv_list[4][1:]
        print(leadership_experience_sublists)
 
    if formatted_cv_list[5] == 0:
        skills_sublists = "ERROR"
    else:
        skills_sublists = formatted_cv_list[5][1:]
        print(skills_sublists)

    # Create Document Object that will then be Populated
    formatted_cv = Document()

    # SECTION 1A - Add Name Header
    section_1a = formatted_cv.add_paragraph()
    run1a = section_1a.add_run(name_header)
    run1a.bold = True
    run1a.font.size = Pt(22)
    run1a.font.name = 'Times New Roman'
    section_1a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    section_1a.space_before = Pt(12)
    section_1a.space_after = Pt(6)

    # SECTION 1B - Add Description Header
    section_1b = formatted_cv.add_paragraph()
    run1b = section_1b.add_run(desc_header[0] + " | " + desc_header[1] + " | " + desc_header[2] + " | " + desc_header[3])
    run1b.bold = False
    run1b.font.size = Pt(11)
    run1b.font.name = 'Times New Roman'
    section_1b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    section_1b.space_before = Pt(12)
    section_1b.space_after = Pt(6)

    # SECTION 2 - Add Education Section 60 spaces
    # Add Header
    section_2_header = formatted_cv.add_paragraph()
    run2_header = section_2_header.add_run("EDUCATION")
    run2_header.bold = True
    run2_header.font.size = Pt(12)
    run2_header.font.name = 'Times New Roman'

    # Add University and Location, Course and Dates Headers
    for sublist in education_section_sublists:
        if sublist[0] == "ERROR":
            run2_univ_loc = section_2_header.add_run("This Section is Empty, Please Format it Accordingly!")  
        else:
            run2_univ_loc = section_2_header.add_run(sublist[0] + " "*60 + sublist[1]) # 60 Spaces, will need to be Manually Fixed on Resultant Document
            run2_course_date = section_2_header.add_run(sublist[2] + " "*60 + sublist[3])

            # Add Description Points of Each Education Entry
            for point in sublist[4]:
                 run2_desc = section_2_header.add_run(point)
    
        # Univeristy and Location Formatting
        run2_univ_loc.bold = True
        run2_univ_loc.font.size = Pt(11)
        run2_univ_loc.font.name = 'Times New Roman'

        # Course and Dates Formatting
        run2_course_date.bold = False
        run2_course_date.italic = True
        run2_course_date.font.size = Pt(11)
        run2_course_date.font.name = 'Times New Roman'

        # Content Formatting
        run2_desc.bold = False
        run2_desc.italic = False
        run2_desc.font.size = Pt(11)
        run2_desc.font.name = 'Times New Roman'

    # SECTION 3 - Add Work Experience Section
    # Add Header
    section_3_header = formatted_cv.add_paragraph()
    run3_header = section_3_header.add_run("WORK EXPERIENCE")
    run3_header.bold = True
    run3_header.font.size = Pt(12)
    run3_header.font.name = 'Times New Roman'

    # Add Company and Location, Position and Dates Headers
    for sublist in work_experience_sublists:
        if sublist[0] == "ERROR":
            run3_comp_loc = section_3_header.add_run("This Section is Empty, Please Format it Accordingly!")  
        else:
            run3_comp_loc = section_3_header.add_run(sublist[0] + " "*60 + sublist[1]) # 60 Spaces, will need to be Manually Fixed on Resultant Document
            run3_pos_date = section_3_header.add_run(sublist[2] + " "*60 + sublist[3])

            # Add Description Points of Each Education Entry
            for point in sublist[4]:
                 run3_desc = section_3_header.add_run(point)
    
        # Company and Location Formatting
        run3_comp_loc.bold = True
        run3_comp_loc.font.size = Pt(11)
        run3_comp_loc.font.name = 'Times New Roman'

        # Position and Dates Formatting
        run3_pos_date.bold = False
        run3_pos_date.italic = True
        run3_pos_date.font.size = Pt(11)
        run3_pos_date.font.name = 'Times New Roman'

        # Content Formatting
        run3_desc.bold = False
        run3_desc.italic = False
        run3_desc.font.size = Pt(11)
        run3_desc.font.name = 'Times New Roman'

    # SECTION 4 - Add Leadership Experience Section
    # Add Header
    section_4_header = formatted_cv.add_paragraph()
    run4_header = section_4_header.add_run("LEADERSHIP EXPERIENCE")
    run4_header.bold = True
    run4_header.font.size = Pt(12)
    run4_header.font.name = 'Times New Roman'

    # Add Company and Location, Position and Dates Headers
    for sublist in leadership_experience_sublists:
        if sublist[0] == "ERROR":
            run4_comp_loc = section_4_header.add_run("This Section is Empty, Please Format it Accordingly!")  
        else:
            run4_comp_loc = section_4_header.add_run(sublist[0] + " "*60 + sublist[1]) # 60 Spaces, will need to be Manually Fixed on Resultant Document
            run4_pos_date = section_4_header.add_run(sublist[2] + " "*60 + sublist[3])

            # Add Description Points of Each Leadership Experience Entry

            for point in sublist[4]:
                 run4_desc = section_4_header.add_run(point)
    
        # Company and Location Formatting
        run4_comp_loc.bold = True
        run4_comp_loc.font.size = Pt(11)
        run4_comp_loc.font.name = 'Times New Roman'

        # Position and Dates Formatting
        run4_pos_date.bold = False
        run4_pos_date.italic = True
        run4_pos_date.font.size = Pt(11)
        run4_pos_date.font.name = 'Times New Roman'

        # Content Formatting
        run4_desc.bold = False
        run4_desc.italic = False
        run4_desc.font.size = Pt(11)
        run4_desc.font.name = 'Times New Roman'

    # SECTION 5 - Add Skills and Additional Training Section
    # Add Header
    section_5_header = formatted_cv.add_paragraph()
    run5_header = section_5_header.add_run("Skills and Additional Training")
    run5_header.bold = True
    run5_header.font.size = Pt(12)
    run5_header.font.name = 'Times New Roman'

    # Add Company and Location, Position and Dates Headers
    for entry in skills_sublists:
        if entry == "ERROR":
            run5_desc = section_5_header.add_run("This Section is Empty, Please Format it Accordingly!")  
        else:
            run5_desc = section_5_header.add_run(entry)

        # Content Formatting
        run5_desc.bold = False
        run5_desc.font.size = Pt(11)
        run5_desc.font.name = 'Times New Roman'

    # Set Single-Line Spacing
    set_line_spacing(formatted_cv, 1)

    formatted_cv.save("test.docx")

def extract_text(file_path):
    _, extension = os.path.splitext(file_path)

    # Deal with PDF Files
    if extension == '.pdf':
        pdf_reader = PdfReader('file_path')

        num_pages = pdf_reader.getNumPages()

        text = ""

        for x in range(num_pages):
            page = pdf_reader.pages[0]

            text = text + " " + pages.extract_text()
        pass

    # Deal with Word/Docx Files
    elif extension == '.docx':
        return process(file_path)
    else:
        # Handle other formats if needed
        pass

    return text

def set_line_spacing(doc, spacing):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(spacing)
        paragraph.space_after = Pt(spacing)

unformatted_cv = """Khamis Mohammad Hamad Almazrouei 
Liwa, Al Dhafra Region 
+971501551120 | khamisliwa9@gmail.com  


CAREER OBJECTIVE:

An innovative UAE National, completed Applied bachelor of Science in Business Administration by JAN 2022 looking for opportunities to develop my career as successful Business man in a high teach and innovative organization, skills and career development opportunities, and to place my efforts with the Employer organization to achieve their strategic objectives.

SUMMARY
 
	Tadbeer\Customer service | 2020-2022 
	Successfully Completed Toefl Exam with 513 as a score | 2018
	First Aid Certification - 2015 
	Abu Dhabi Sewerage Service company | 2014-2015
	Attended workshop(Mastering the interview skills and follow-up strategy)-2014 
	Professional Experiences Training Center-2013
	Did several volunteer activities

Technical skills & Soft Skills
	I have good skills in the MS-Office tools which includes MS-Word, MS-Excel and MS-Power point.
	I have good communication skills that helps me to make friends with the unknown people and gain new relation ship
	 I have developed good skills in facing the challenges of the real-world and finds the solutions to solve them.
	 Confident, articulate in dealing with tasks
	I demonstrate good teamwork
	I am capable of looking for innovative solutions to the problems and I am highly self-motivated
	Good in management of time
	Arabic as native, good in English( Speak, Write, understand) and can speak Hindi language

EDUCATION

2017 – 2022 | Institute of Management Technology, Dubai
Applied Bachelor Science in Business Administration
GPA 2.99
Major Courses:
	Global Logistics Management
	Advance Managerial Accounting 
2013 | Farouq School,Liwa
Higher School Certificate
Marks : 79 %


INTERNSHIPS / VOLUNTEERING

Volunteering activities

Organization Name	Days / Month/Year	Activities 
Redcresent, Madinat Zayed	1 day
March 2017	Did filing, data feeding 
Ministry of culture and Knowledge development- Madinat Zayed	2 weeks, July 2017	Well-handled emails and phone calls for the summer cultural event
Ministry of culture and Knowledge development- Madinat Zayed	1 Week
August 2018	Printing and files arrangements
Emirates Foundation Forum	One day
Sep 2018	Printing and files arrangements
ADSSC-Abu Dhabi	1 year 2014-2015	Customer survives 


Interests and Hobbies

	Book reading | Volunteer activities |taking time to work on own business projects at home 

References

	Mr. Rashia Mohammed - Officer- Work Placement – Institute of Management Technology, rasha@inmt.ac.ae , 056 7573999	
"""

formatted_cv_list = cv_reader(unformatted_cv)

docx_writer(formatted_cv_list)