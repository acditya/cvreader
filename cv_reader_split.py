import ast
import os
import openai
from docx2txt import process
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Input and Output Folder Global Definitions
cv_folder = "CVs"
output_folder = "Formatted CVs"

def cv_reader(unformatted_cv):
    # Set OpenAI API key & Initialize GPT Client
    # api_key part 1 is sk-zSGM2JQSdiQ
    # api_key part 2 is ChxnLDfvFT3BlbkFJJrNgy8gHCjscsu1OJJTU
    openai.api_key = 'PLACEHOLDER'

    # Initialize AI's Personality
    conversation = [
    {"role": "system", "content": """You are a CV-conversion assistant whose job is to take a CV and standardize it in a given format and return it to me in a part-by-part sequential manner. I will pass you an entire CV and you will understand it in whole, then I will pass you the specifics I'd like you to give me."""},
    ]
    conversation.append({"role": "assistant", "content": "Certainly! Please provide the CV you'd like me to convert, and let me know the specific format or information you want in each part. Once you've shared the CV, you can also specify how you'd like me to organize and present the information."})
    
    # Unformatted CV Input
    conversation.append({"role": "user", "content": unformatted_cv})
    
    # Sample Hardcoded Output
    conversation.append({"role": "assistant", "content": "Thank you for providing the CV. Please let me know the specific format or information you would like to receive, and I'll organize the details accordingly. For example, you might want the CV in sections such as Personal Information, Career Objective, Summary, Technical & Soft Skills, Education, Internships/Volunteering, Interests and Hobbies, and References. Additionally, if there are specific details you want to highlight or exclude, please specify"})
    
    conversation.append({"role": "user", "content": "I now want you to pass me an array for each section, and nothing else except for it, not even a single additional character except for the list. You should return nothing except the array, no text nor any additional commentary. Now, only give me the array for first section. The first section should have 5 elements in the following format [Name, Location, LinkedIn, Phone, Email]; if one or more of these elements are missing, leave the placeholder there such as LINKEDIN_MISSING or NAME_MISSING such that the array is still 5 elements long."})

    personal_information = chat_with_model(conversation)
    conversation.append({"role": "assistant", "content": personal_information})

    conversation.append({"role": "user", "content": "Now pass me only the second section array which should contain all educational experiences and nothing else except for it, not even a single additional character except for the list. You should return nothing except the array, no text nor any additional commentary. Each educational experience should be a subarray of length 5: [University/School, Location, Major/Course, Dates,  Points such as GPA, Organizations, Coursework]. If any of these are missing, kindly leave a placeholder as you did before such as LOCATION_PLACEHOLDER or MAJOR/COURSE_PLACEHOLDER. Ultimately, you will return me an array of these subsection arrays, there should be one subsection array for each individual educational experience. Return nothing except for the array of arrays, no text nor any additional commentary."})

    educational_information = chat_with_model(conversation)
    conversation.append({"role": "assistant", "content": educational_information})

    conversation.append({"role": "user", "content": "Now pass me only the third section which should contain all work, internship, and major relevant volunteering experiences and nothing else except for it, not even a single additional character except for the list. You should return nothing except the array, no text nor any additional commentary. Each work experience should be an array of length 5: [Company/Organization, Location, Position/Role, Dates,  description_subarray]. The description subarray will be used to describe each experience in a bulleted format, each entry in the description subarray should contain an explanation of what was done in the role, how it was done and most importantly if applicable, any notable achievements. If any of these are missing, kindly leave a placeholder as you did before such as LOCATION_PLACEHOLDER or [DESCRIPTION_SUBARRAY_PLACEHOLDER] for the description_subarray. Ultimately, you will return me an array of these subsection arrays, there should be one subsection array for each individual work experience. Return nothing except for the array of arrays, no text nor any additional commentary."})

    work_information = chat_with_model(conversation)
    conversation.append({"role": "assistant", "content": work_information})

    conversation.append({"role": "user", "content": "Now pass me only the fourth section with should contain all skills and additional training. Return this in a paragraph format that outlines each skill/training entry where successive entries are separated by commas. If there are no skills to be found on the original CV, first try your best to create some based on what you know about the CV. If this is not feasible, simply return a SKILLS_PLACEHOLDER. Remember, this last fourth section needs only a paragraph string, not a list."})
    
    skills_information = chat_with_model(conversation)
    conversation.append({"role": "assistant", "content": skills_information})

    # Deal with List/Non-List of this Section due to ChatGPT Unpredictability
    if skills_information is list:
        return ([ast.literal_eval(personal_information), ast.literal_eval(educational_information), ast.literal_eval(work_information), ast.literal_eval(skills_information)])
    else:
        return ([ast.literal_eval(personal_information), ast.literal_eval(educational_information), ast.literal_eval(work_information), skills_information])


# Helper Function to Create a conversation and get AI-generated responses
def chat_with_model(messages):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    return response.choices[0].message["content"]

def docx_writer(formatted_cv_list):
    # Takes a Formatted List and Outputs Desired CV as per Template
    # List [PERSONAL INFO, EDUCATIONAL INFO, WORK INFO, SKILLS INFO]

    p_info = formatted_cv_list[0]
    # [Name, Location, LinkedIn, Phone, Email]
    e_info = formatted_cv_list[1]
    # List of sublists such as [Uni/School, Location, Major/Course, Dates, Description]
    w_info = formatted_cv_list[2]
    # List of sublists such as [Company, Location, Position, Dates, [SUBLIST OF BULLET-POINT DESCRIPTION]]
    s_info = formatted_cv_list[3]
    # String/Paragraph

    # Create Document Object that will then be Populated
    formatted_cv = Document()

    # Add Name Header
    section_1a = formatted_cv.add_paragraph()
    run1a = section_1a.add_run(p_info[0])
    run1a.bold = True
    run1a.font.size = Pt(22)
    run1a.font.name = 'Times New Roman'
    section_1a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    section_1a.space_before = Pt(12)
    section_1a.space_after = Pt(6)

    # Add Rest of Personal Info
    section_1b = formatted_cv.add_paragraph()
    run1b = section_1b.add_run(p_info[1] + " | " + p_info[2] + " | " + p_info[3] + " | " + p_info[4])
    run1b.bold = False
    run1b.font.size = Pt(11)
    run1b.font.name = 'Times New Roman'
    section_1b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    section_1b.space_before = Pt(12)
    section_1b.space_after = Pt(6)

    # Add Education Section 60 spaces
    # Add Header
    educ = formatted_cv.add_paragraph()
    e_header = educ.add_run("Education")
    e_header.bold = True
    e_header.font.size = Pt(12)
    e_header.font.name = 'Times New Roman'
    educ.add_run().add_break()

    # Add University and Location, Course and Dates Headers
    e_content = formatted_cv.add_paragraph()
    e_content.style = 'List Bullet'
    for entry in e_info:
        e_univ_loc = educ.add_run(entry[0] + " "*25 + entry[1]) # 25 Spaces, will need to be Manually Fixed on Resultant Document
        educ.add_run().add_break()
        e_course_date = educ.add_run(entry[2] + " "*25 + entry[3]) # 25 Spaces, will need to be Manually Fixed on Resultant Document
        educ.add_run().add_break()

        # Add Description of Each Education Entry
        e_desc = e_content.add_run(entry[4])
            # section_2_content.add_run().add_break()
        
        # Univeristy and Location Formatting
        if e_univ_loc is not None:
            e_univ_loc.bold = True
            e_univ_loc.font.size = Pt(11)
            e_univ_loc.font.name = 'Times New Roman'

        # Course and Dates Formatting
        if e_course_date is not None:
            e_course_date.bold = False
            e_course_date.italic = True
            e_course_date.font.size = Pt(11)
            e_course_date.font.name = 'Times New Roman'

        # Content Formatting
        if e_desc is not None:
            e_desc.bold = False
            e_desc.italic = False
            e_desc.font.size = Pt(11)
            e_desc.font.name = 'Times New Roman'

    # Add Work Experience Section
    # Add Header
    work_exp = formatted_cv.add_paragraph()
    w_header = work_exp.add_run("Work Experience")
    w_header.bold = True
    w_header.font.size = Pt(12)
    w_header.font.name = 'Times New Roman'
    work_exp.add_run().add_break()

    # Add Company and Location, Position and Dates Headers
    w_content = formatted_cv.add_paragraph()
    w_content.style = 'List Bullet'
    for entry in w_info:
        w_comp_loc = work_exp.add_run(entry[0] + " "*25 + entry[1]) # 25 Spaces, will need to be Manually Fixed on Resultant Document
        work_exp.add_run().add_break()
        w_pos_date = work_exp.add_run(entry[2] + " "*25 + entry[3]) # 25 Spaces, will need to be Manually Fixed on Resultant Document
        work_exp.add_run().add_break()

        # Add Description Points of Each Work Experience Entry
        for point in entry[4]:
            w_desc = w_content.add_run(point)
            # section_3_content.add_run().add_break()
    
        # Company and Location Formatting
        if w_comp_loc is not None:
            w_comp_loc.bold = True
            w_comp_loc.font.size = Pt(11)
            w_comp_loc.font.name = 'Times New Roman'

        # Position and Dates Formatting
        if w_pos_date is not None:
            w_pos_date.bold = False
            w_pos_date.italic = True
            w_pos_date.font.size = Pt(11)
            w_pos_date.font.name = 'Times New Roman'

        # Content Formatting
        if w_desc is not None:
            w_desc.bold = False
            w_desc.italic = False
            w_desc.font.size = Pt(11)
            w_desc.font.name = 'Times New Roman'

    # Add Skills and Additional Training Section
    # Add Header
    skills = formatted_cv.add_paragraph()
    s_header = skills.add_run("Skills and Additional Training")
    s_header.bold = True
    s_header.font.size = Pt(12)
    s_header.font.name = 'Times New Roman'

    # Add Content
    s_content = formatted_cv.add_paragraph()
    if s_info is not list: # Deal with List/Non-List of this Section due to ChatGPT Unpredictability
        s_content = skills.add_run(s_info)
    else:
        s_content = skills.add_run(s_info[0])

        # Content Formatting
        if s_content is not None:
            s_content.bold = False
            s_content.font.size = Pt(11)
            s_content.font.name = 'Times New Roman'

    # Save the Document
    formatted_cv.save("modular-test.docx")

def extract_text(file_path):
    _, extension = os.path.splitext(file_path)

    # Deal with PDF Files
    if extension == '.pdf':
        pdf_reader = PdfReader('file_path')

        num_pages = pdf_reader.getNumPages()

        text = ""

        for x in range(num_pages):
            page = pdf_reader.pages[0]

            text = text + " " + pages.extract_text()
        pass

    # Deal with Word/Docx Files
    elif extension == '.docx':
        return process(file_path)
    else:
        # Handle other formats if needed
        pass

    return text


unformatted_cv = """
		Anwaar Aljaafari
 
PROFILE
I am a highly motivated and enthusiastic person who would thrive in a challenging environment. I am work well under pressure and am a problem solver. I take the initiative and enjoy working as part of a team. I’m interested in opportunities in Marketing and Communications, Public Relations, as well as EHS opportunities, either in Al Ain or Abu Dhabi City.


CONTACT
PHONE:
050-962-0027 / 050-833-6736  
 

EMAIL: anown.93@gmail.com

HOBBIES
Exploring new technologies. 
Interior designing 
Like meeting new people.
Reading 






		
EDUCATION 
KHAWARIZMI INTERNATIONAL COLLEGE, 2018-2021   
Bachelor of Mass Commutation (Digital Media),AGPA (3.69)

ABU DHABI VOCATIONAL EDUCATION AND TRAINING INSTITUTE. 2013- 2015 
Higher Diploma Environment Health and Safety, AGPA (3.46)

WORK EXPERIENCE
Silal for Food & Technology, Officer Packing Material 
2021 – 2023
•	Preparing files and an internal network that ensures input and output information, and quantities of packaging materials received and sent to suppliers are documented.
•	Working on checking the quantities of available packing materials.
•	Planning in advance to ensure needs are met in times of crisis. 

 Abu Dhabi Authority for Agriculture and Food Safety, Inventory Supervisor 
2016 – 2021

•	Conducted inspection, quality, and warehouse management activities
•	Inspected products and their conformity with the approved specifications
SKILLS
Advanced level Microsoft Office 

Run Social Media Campaigns

•	Video Preparation & Poster Design

Language Skills 
Arabic: 
Fluent in written and spoken (native speaker),  

English: 
Excellent reading, writing, and speaking.
	•	The ability to handle difficult situations.
•	Audit and Inspection in Health and Safety
•	Manage and organize events.
•	Listening and following instructions.
•	Interpersonal skills 
•	Familiarity with the scientific origins and foundations of public relations and the methods of analyzing and measuring public opinion.
"""

formatted_cv_list = cv_reader(unformatted_cv)

docx_writer(formatted_cv_list)